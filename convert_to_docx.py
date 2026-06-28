import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# MD 파일 읽기
with open('D:\\result\\korean_football_history.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# DOCX 문서 생성
doc = Document()

# 스타일 설정
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(11)

# 제목 추가
title = doc.add_heading('한국 축구 역사', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()  # 여백

# 목차 생성
doc.add_heading('목차', level=1)
sections = re.findall(r'## \d+\. (.+?)(?=\n##|\n##|\Z)', md_content)
for i, section in enumerate(sections, 1):
    clean_title = re.sub(r'\*\*', '', section)
    p = doc.add_paragraph()
    p.add_run(f"{i}. {clean_title}").italic = True

doc.add_page_break()

# 본문 파싱 및 추가
lines = md_content.split('\n')
current_section = None
current_level = 0

for line in lines:
    line = line.strip()
    
    # 섹션 헤딩
    if line.startswith('## '):
        section_text = line[3:].strip()
        # 섹션 번호와 제목 분리
        match = re.match(r'(\d+)\.\s*(.+)', section_text)
        if match:
            level = int(match.group(1))
            title_text = match.group(2)
            clean_title = re.sub(r'\*\*', '', title_text)
            # 이미지 placeholder 추가
            doc.add_heading(f"{match.group(1)}. {clean_title}", level=1)
            p = doc.add_paragraph()
            p.add_run(f"[이미지: {clean_title} 관련 사진]").italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
    
    # 표 처리
    elif '| 연도 | 대회 | 성적 | 주요 내용 |' in line or '|------|' in line or '| 19' in line:
        continue  # 표는 별도 처리
    # 표 데이터 행 건너뛰기 (표가 있는 섹션의 일반 텍스트)
    elif line.startswith('| ') and '---' not in line and '|' in line:
        continue
    
    # 일반 텍스트 (굵은 글씨 처리)
    elif line and not line.startswith('#') and not line.startswith('>') and not line.startswith('[^'):
        # ** 굵은 글씨 처리
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        # [^숫자] 참고문헌 표시 제거
        text = re.sub(r'\[\^[0-9]+\]', '', text)
        if text.strip():
            doc.add_paragraph(text)

doc.add_page_break()

# 참고문헌
doc.add_heading('참고 문헌', level=1)
references = re.findall(r'\[\^[0-9]+\]:\s*(.+)', md_content)
for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.5)

# 결재 라인
doc.add_paragraph()
doc.add_paragraph()
sign_block = doc.add_paragraph()
sign_block.add_run('결 재 라 인').bold = True
sign_block.alignment = WD_ALIGN_PARAGRAPH.RIGHT

sign_table = doc.add_table(rows=3, cols=4)
sign_table.style = 'Table Grid'
cells = sign_table._cells
labels = ['작성자', '검토자', '승인자', '최종결재']
for i, label in enumerate(labels):
    cells[i].text = f'{label}\n\n\n\n(서명)　　(날짜)'

# 파일 저장
doc.save('D:\\result\\korean_football_history.docx')
print("DOCX 파일 생성 완료: D:\\result\\korean_football_history.docx")